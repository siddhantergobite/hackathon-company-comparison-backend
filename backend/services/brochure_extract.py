"""
Brochure / company profile extraction for pitching company (Exhibit A).
Supports PDF, DOCX, images (vision), and company-name website search.
"""
from __future__ import annotations

import io
import json
import re
from typing import Optional

from backend.services import llm as llm_client

MODEL = llm_client.AZURE_MODEL


def _extract_pdf(data: bytes) -> tuple[str, int]:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=data, filetype="pdf")
        pages = []
        for page in doc:
            pages.append(page.get_text())
        text = "\n".join(pages).strip()
        return text, len(doc)
    except ImportError:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        pages = [p.extract_text() or "" for p in reader.pages]
        text = "\n".join(pages).strip()
        return text, len(reader.pages)


def _extract_docx(data: bytes) -> tuple[str, int]:
    from docx import Document
    doc = Document(io.BytesIO(data))
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            paras.extend(c.text.strip() for c in row.cells if c.text.strip())
    text = "\n".join(paras).strip()
    word_count = len(text.split())
    return text, max(1, word_count // 300)


def _extract_image(data: bytes) -> tuple[str, int]:
    import base64
    from PIL import Image
    img = Image.open(io.BytesIO(data)).convert("RGB")
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=85)
    b64 = base64.b64encode(buf.getvalue()).decode()
    text = llm_client.chat_vision(
        "Extract ALL text from this brochure/page exactly. Return plain text only, preserve structure.",
        b64,
        max_tokens=4000,
    )
    return text.strip(), 1


def _extract_text(data: bytes, filename: str) -> tuple[str, int, str]:
    ext = (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()
    if ext == "pdf":
        text, pages = _extract_pdf(data)
        return text, pages, "pdf"
    if ext in ("docx", "doc"):
        text, pages = _extract_docx(data)
        return text, pages, "docx"
    if ext in ("png", "jpg", "jpeg", "webp", "gif", "bmp"):
        text, pages = _extract_image(data)
        return text, pages, "image"
    # fallback: try utf-8 text
    try:
        text = data.decode("utf-8", errors="ignore").strip()
        return text, 1, "text"
    except Exception:
        raise ValueError(f"Unsupported file type: .{ext or 'unknown'}")


def _structure_with_llm(raw_text: str, source_hint: str = "") -> dict:
    trimmed = raw_text[:12000]
    word_count = len(raw_text.split())
    prompt = f"""You are a B2B sales intelligence analyst. Extract a structured company profile from this brochure/company material.

SOURCE: {source_hint or "Uploaded document"}
WORD COUNT: {word_count}

IMPORTANT: Extract information ONLY about the company described in the source text.
If the text is unrelated or about a different company, return empty services/industries and say so in summary.

TEXT:
{trimmed}

Return ONLY valid JSON with this exact schema:
{{
  "company_name": "string",
  "summary": "2-3 sentence company overview",
  "services": ["service 1", "service 2"],
  "industries": ["industry 1"],
  "case_studies": "notable achievements / client wins as one paragraph",
  "contacts": [{{"name": "...", "title": "...", "email": "...", "phone": "..."}}],
  "website": "url if found or empty string",
  "verbatim_extract": "200-400 word representative excerpt from the source text"
}}

CONTACT RULES: Only list founders, directors, C-suite (CEO/MD/CTO), or official company inbox in contacts.
Do NOT list delivery heads, sales reps, engineers, or department staff as primary contact.
If MCA/registry directors are mentioned in the text, use those first."""
    raw = llm_client.chat(
        [{"role": "user", "content": prompt}],
        temperature=0.2,
        max_tokens=2500,
        json_mode=True,
    )
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        m = re.search(r"\{[\s\S]*\}", raw)
        if m:
            return json.loads(m.group())
        raise RuntimeError("Failed to parse brochure extraction")


_LEADERSHIP_TITLE = re.compile(
    r"founder|co-founder|cofounder|ceo|chief|director|president|owner|chairman|managing|partner|proprietor|md\b",
    re.I,
)
_STAFF_TITLE = re.compile(
    r"delivery|sales|engineer|developer|consultant|support|representative|coordinator|analyst|specialist|head of delivery",
    re.I,
)


def _director_email_match(director_name: str, email: str) -> bool:
    if not director_name or not email:
        return False
    first = director_name.split()[0].lower()
    local = email.split("@")[0].lower()
    return first in local.replace(".", " ").replace("_", " ")


def _contacts_from_registry(scraped: dict) -> list:
    """MCA directors from ZaubaCorp — highest-confidence contacts for Indian cos."""
    conf = scraped.get("zauba_match_confidence") or "Low"
    if conf == "Low":
        return []
    zauba = scraped.get("zaubacorp_structured") or {}
    directors = zauba.get("Directors") or []
    if not directors:
        return []
    reg_email = (zauba.get("Email Address") or "").strip().lower()
    contacts = []
    for i, d in enumerate(directors):
        if not isinstance(d, dict) or not d.get("name"):
            continue
        name = d["name"]
        email = reg_email if (reg_email and (i == 0 or _director_email_match(name, reg_email))) else ""
        contacts.append({
            "name": name,
            "title": d.get("designation") or "Director",
            "email": email,
            "phone": "",
            "source": "ZaubaCorp MCA",
            "confidence": conf,
        })
    return contacts


def _contact_rank(c: dict) -> int:
    source = (c.get("source") or "").lower()
    title = c.get("title") or ""
    if "zaubacorp" in source or "mca" in source:
        return 0
    if _LEADERSHIP_TITLE.search(title):
        return 1
    if _STAFF_TITLE.search(title) and not _LEADERSHIP_TITLE.search(title):
        return 3
    return 2


def _sort_contacts(contacts: list) -> list:
    return sorted(contacts, key=_contact_rank)[:8]


def _contacts_from_scrape(scraped: dict) -> list:
    """Pull named phones/emails from website contact scrape."""
    cd = scraped.get("contact_data") or {}
    contacts = []
    seen = set()
    for p in cd.get("phones") or []:
        num = p.get("number") or ""
        if not num or num in seen:
            continue
        seen.add(num)
        contacts.append({
            "name": p.get("person_name") or p.get("name") or "",
            "title": p.get("title") or p.get("role") or "Business Contact",
            "email": "",
            "phone": num,
        })
    for e in cd.get("emails") or []:
        em = e.get("email") or ""
        if not em:
            continue
        name = e.get("person_name") or e.get("name") or ""
        if not name and "@" in em:
            local = em.split("@")[0].replace(".", " ").replace("_", " ")
            name = local.title()
        contacts.append({
            "name": name,
            "title": e.get("title") or e.get("label") or e.get("role") or "",
            "email": em,
            "phone": "",
            "source": "Company Website",
            "confidence": "Medium",
        })
    return _sort_contacts(contacts)


def _merge_profile_contacts(profile: dict, scraped: dict) -> dict:
    registry = _contacts_from_registry(scraped)
    scraped_contacts = _contacts_from_scrape(scraped)
    llm_contacts = profile.get("contacts") or []

    merged = []
    seen_emails, seen_names = set(), set()

    def _add(c):
        if not c or not (c.get("name") or c.get("email")):
            return
        key = (c.get("email") or "").lower(), (c.get("name") or "").lower()
        if key[0] in seen_emails or (key[1] and key[1] in seen_names):
            return
        if key[0]:
            seen_emails.add(key[0])
        if key[1]:
            seen_names.add(key[1])
        merged.append(c)

    for c in registry:
        _add(c)
    for c in llm_contacts:
        if isinstance(c, dict):
            title = c.get("title") or ""
            if registry and _STAFF_TITLE.search(title) and not _LEADERSHIP_TITLE.search(title):
                continue
            _add({**c, "source": c.get("source") or "Brochure/Website", "confidence": c.get("confidence") or "Medium"})
    for c in scraped_contacts:
        if registry and _contact_rank(c) >= 3:
            continue
        _add(c)

    profile["contacts"] = _sort_contacts(merged) if merged else llm_contacts
    if registry:
        profile["registry_intelligence"] = {
            "source": "ZaubaCorp MCA",
            "url": scraped.get("zaubacorp_url") or "",
            "directors": [{"name": d.get("name"), "designation": d.get("designation")} for d in registry],
        }
    return profile


def from_file(data: bytes, filename: str) -> dict:
    text, pages, file_type = _extract_text(data, filename)
    if not text or len(text.strip()) < 30:
        raise ValueError("Could not extract enough text from the uploaded file.")
    profile = _structure_with_llm(text, f"Uploaded {file_type}: {filename}")
    profile["_meta"] = {
        "source": "upload",
        "filename": filename,
        "file_type": file_type,
        "pages": pages,
        "word_count": len(text.split()),
        "model_used": MODEL,
    }
    return profile


def _is_url(text: str) -> bool:
    t = text.strip()
    if t.startswith(("http://", "https://")):
        return True
    return bool(re.match(r"^[a-zA-Z0-9][-a-zA-Z0-9.]*\.[a-zA-Z]{2,}(/.*)?$", t))


def _normalize_url(text: str) -> str:
    t = text.strip()
    if not t.startswith(("http://", "https://")):
        t = "https://" + t
    return t


def _collect_scrape_text(scraped: dict) -> str:
    """Merge homepage + all sub-pages into one text blob."""
    parts = [
        scraped.get("title", ""),
        scraped.get("description", ""),
        scraped.get("keywords", ""),
        scraped.get("homepage_text", ""),
        scraped.get("about_text", ""),
        scraped.get("products_text", ""),
        scraped.get("pricing_text", ""),
        scraped.get("careers_text", ""),
        scraped.get("leadership_text", ""),
        scraped.get("blog_text", ""),
    ]
    return "\n\n".join(p.strip() for p in parts if p and p.strip())


def _report_to_brochure(report: dict, url: str) -> dict:
    """Map full company_research report → brochure profile (same quality as Exhibit B)."""
    co = report.get("company_profile") or {}
    meta = report.get("_meta") or {}
    prod = report.get("products_services") or {}
    reg = report.get("registry_intelligence") or {}
    leaders = report.get("leadership_team") or []
    contacts_raw = (report.get("contact_intelligence") or {}).get("emails") or []

    services = []
    for item in prod.get("primary_offerings") or []:
        if isinstance(item, dict):
            v = item.get("item") or item.get("value") or item.get("name")
            if v:
                services.append(str(v))
        elif isinstance(item, str):
            services.append(item)

    contacts = []
    reg_email = (reg.get("email") or "").strip().lower()
    for ldr in leaders:
        if not isinstance(ldr, dict) or not ldr.get("name"):
            continue
        name = ldr["name"]
        email = reg_email if _director_email_match(name, reg_email) else ""
        if not email:
            for e in contacts_raw:
                if isinstance(e, dict) and _director_email_match(name, e.get("email", "")):
                    email = e.get("email", "")
                    break
        contacts.append({
            "name": name,
            "title": ldr.get("role") or "Director",
            "email": email,
            "phone": "",
            "source": "ZaubaCorp MCA",
            "confidence": "High",
        })

    if not contacts:
        for e in contacts_raw[:5]:
            if isinstance(e, dict):
                title = e.get("title") or e.get("role") or ""
                if _STAFF_TITLE.search(title) and not _LEADERSHIP_TITLE.search(title):
                    continue
                contacts.append({
                    "name": e.get("person_name") or e.get("name") or "",
                    "title": title,
                    "email": e.get("email") or "",
                    "source": e.get("source") or "Public Web",
                    "confidence": e.get("confidence") or "Medium",
                })

    contacts = _sort_contacts(contacts)

    def _v(field):
        if isinstance(field, dict):
            return field.get("value") or ""
        return str(field or "")

    name = meta.get("company_name") or co.get("name") or ""
    summary = _v(co.get("description")) or report.get("intelligence_score", {}).get("summary", "")

    return {
        "company_name": name,
        "summary": summary,
        "services": services or [s for s in (_v(co.get("industry")),) if s],
        "industries": [_v(co.get("industry"))] if _v(co.get("industry")) else [],
        "case_studies": _v(report.get("content_strategy", {}).get("key_themes")) if isinstance(report.get("content_strategy"), dict) else "",
        "contacts": contacts,
        "website": url,
        "registry_intelligence": reg if reg else None,
        "verbatim_extract": summary[:600],
        "_meta": {
            "source": "url_research",
            "website_url": url,
            "word_count": len(summary.split()),
            "model_used": MODEL,
            "pipeline": "company_research",
        },
    }


def _pick_company_url(company_name: str, results: list) -> str:
    """Pick the best official website from search results."""
    from urllib.parse import urlparse

    slug = re.sub(r"[^a-z0-9]", "", company_name.lower())
    blocked = {
        "microsoft.com", "google.com", "facebook.com", "twitter.com", "x.com",
        "youtube.com", "wikipedia.org", "linkedin.com", "instagram.com",
        "apple.com", "amazon.com", "bing.com", "yahoo.com",
    }
    candidates = []
    for r in results:
        href = r.get("href", "")
        if not href.startswith("http"):
            continue
        domain = urlparse(href).netloc.replace("www.", "").lower()
        if any(b in domain for b in blocked):
            continue
        score = 0
        if slug and slug in domain.replace(".", "").replace("-", ""):
            score += 10
        if slug and slug in href.lower():
            score += 5
        title = (r.get("title") or "").lower()
        if company_name.lower() in title:
            score += 3
        if "official" in title or "home" in title:
            score += 1
        candidates.append((score, href, domain))

    if candidates:
        candidates.sort(key=lambda x: -x[0])
        return candidates[0][1]

    # Direct domain guesses
    slug_hyphen = re.sub(r"[^a-z0-9]+", "-", company_name.lower()).strip("-")
    slug_plain = re.sub(r"[^a-z0-9]+", "", company_name.lower())
    for guess in [
        f"https://www.{slug_plain}.com",
        f"https://{slug_plain}.com",
        f"https://www.{slug_hyphen}.com",
        f"https://{slug_plain}.com/us/",
    ]:
        try:
            import requests
            resp = requests.head(guess, timeout=8, allow_redirects=True,
                                 headers={"User-Agent": "Mozilla/5.0"})
            if resp.status_code < 400:
                return resp.url or guess
        except Exception:
            continue
    return ""


def from_company_search(company_name: str) -> dict:
    """Search company website and extract pitching profile."""
    from backend.services import company_research
    query = company_name.strip()
    if not query:
        raise ValueError("Company name or URL is required")

    # ── URL pasted directly → scrape that site (skip DuckDuckGo entirely) ──
    if _is_url(query):
        url = _normalize_url(query)
        scraped = company_research._scrape_website(url)
        combined = _collect_scrape_text(scraped)
        word_count = len(combined.split())

        # If site is JS-heavy and scrape is thin, use full research pipeline
        if word_count < 150:
            report = company_research.run(url)
            return _report_to_brochure(report, url)

        profile = _structure_with_llm(combined, f"Official website: {url}")
        profile = _merge_profile_contacts(profile, scraped)
        profile["website"] = url
        profile["_meta"] = {
            "source": "url",
            "query": query,
            "website_url": url,
            "word_count": word_count,
            "model_used": MODEL,
        }
        return profile

    # ── Company name only → guess domain first, then DuckDuckGo ──
    slug_plain = re.sub(r"[^a-z0-9]+", "", query.lower())
    url = ""
    for guess in [
        f"https://www.{slug_plain}.com",
        f"https://{slug_plain}.com",
        f"https://{slug_plain}.com/us/",
    ]:
        try:
            import requests
            resp = requests.head(guess, timeout=8, allow_redirects=True,
                                 headers={"User-Agent": "Mozilla/5.0"})
            if resp.status_code < 400:
                url = resp.url or guess
                break
        except Exception:
            continue

    results = []
    if not url:
        for q in [
            f"{query} official website",
            f"{query} company homepage",
            f"site:{slug_plain}.com {query}",
        ]:
            results.extend(company_research._ddg_search(q, max_results=5))
        url = _pick_company_url(query, results)

    if url:
        scraped = company_research._scrape_website(url)
        combined = _collect_scrape_text(scraped)
    else:
        combined = company_research._snippets(results, 5000) or query
        scraped = {"title": query}

    word_count = len(combined.split())

    # Reject garbage: if text doesn't mention company at all and is very short
    slug_check = slug_plain[:4] if len(slug_plain) >= 4 else slug_plain
    if slug_check and slug_check not in combined.lower() and word_count < 200:
        if url:
            report = company_research.run(url)
            return _report_to_brochure(report, url)
        raise ValueError(
            f"Could not find reliable public data for '{query}'. "
            "Try pasting the full company URL (e.g. https://ergobite.com/us/) instead."
        )

    if word_count < 150 and url:
        report = company_research.run(url)
        return _report_to_brochure(report, url)

    profile = _structure_with_llm(combined, f"Web scrape of {url or query}")
    if url:
        profile = _merge_profile_contacts(profile, scraped)
    profile["website"] = url or profile.get("website", "")
    profile["_meta"] = {
        "source": "search",
        "query": query,
        "website_url": url or "",
        "word_count": word_count,
        "model_used": MODEL,
    }
    return profile
